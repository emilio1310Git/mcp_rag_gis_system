"""Procesador de documentos mejorado"""

import logging
from pathlib import Path
from typing import Optional
import PyPDF2
import pandas as pd
import docx
from config import settings

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Procesador de documentos con soporte para múltiples formatos"""
    
    def __init__(self):
        self.supported_extensions = settings.rag.supported_extensions
    
    async def process_file(self, file_path: Path) -> Optional[str]:
        """Procesar archivo según su extensión"""
        try:
            extension = file_path.suffix.lower()
            
            if extension == '.md':
                return await self._process_markdown(file_path)
            elif extension == '.pdf':
                return await self._process_pdf(file_path)
            elif extension == '.csv':
                return await self._process_csv(file_path)
            elif extension == '.txt':
                return await self._process_text(file_path)
            elif extension == '.docx':
                return await self._process_docx(file_path)
            else:
                logger.warning(f"Formato no soportado: {extension}")
                return None
                
        except Exception as e:
            logger.error(f"Error procesando {file_path}: {e}")
            return None
    
    async def _process_markdown(self, file_path: Path) -> str:
        """Procesar archivo Markdown"""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        # Limpiar y normalizar contenido
        content = self._clean_text(content)
        return content
    
    async def _process_pdf(self, file_path: Path) -> str:
        """Procesar archivo PDF"""
        text = ""
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Página {page_num + 1} ---\n"
                        text += page_text
                except Exception as e:
                    logger.warning(f"Error extrayendo página {page_num + 1} de {file_path}: {e}")
        
        return self._clean_text(text)
    
    async def _process_csv(self, file_path: Path) -> str:
        """Procesar archivo CSV"""
        try:
            df = pd.read_csv(file_path)
            
            # Crear descripción textual del CSV
            text = f"Archivo CSV: {file_path.name}\n"
            text += f"Número de filas: {len(df)}\n"
            text += f"Número de columnas: {len(df.columns)}\n"
            text += f"Columnas: {', '.join(df.columns.tolist())}\n\n"
            
            # Información de cada columna
            text += "Descripción de columnas:\n"
            for col in df.columns:
                col_info = f"- {col}: "
                
                if df[col].dtype in ['int64', 'float64']:
                    col_info += f"numérica (min: {df[col].min()}, max: {df[col].max()}, promedio: {df[col].mean():.2f})"
                else:
                    unique_count = df[col].nunique()
                    col_info += f"categórica ({unique_count} valores únicos)"
                    
                    if unique_count <= 10:
                        unique_values = df[col].unique()[:5]
                        col_info += f", ejemplos: {', '.join(map(str, unique_values))}"
                
                text += col_info + "\n"
            
            # Primeras filas como ejemplo
            text += f"\nPrimeras {min(5, len(df))} filas:\n"
            text += df.head().to_string()
            
            # Estadísticas descriptivas para columnas numéricas
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                text += "\n\nEstadísticas descriptivas (columnas numéricas):\n"
                text += df[numeric_cols].describe().to_string()
            
            return text
            
        except Exception as e:
            logger.error(f"Error procesando CSV {file_path}: {e}")
            return f"Error procesando archivo CSV: {str(e)}"
    
    async def _process_text(self, file_path: Path) -> str:
        """Procesar archivo de texto plano"""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        return self._clean_text(content)
    
    async def _process_docx(self, file_path: Path) -> str:
        """Procesar archivo Word DOCX"""
        try:
            doc = docx.Document(str(file_path))
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Procesar tablas si existen
            for table in doc.tables:
                text += "\n--- Tabla ---\n"
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    text += row_text + "\n"
                text += "--- Fin Tabla ---\n"
            
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"Error procesando DOCX {file_path}: {e}")
            return f"Error procesando archivo DOCX: {str(e)}"
    
    def _clean_text(self, text: str) -> str:
        """Limpiar y normalizar texto"""
        if not text:
            return ""
        
        # Normalizar espacios en blanco
        text = ' '.join(text.split())
        
        # Eliminar caracteres de control
        text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t')
        
        # Limitar longitud si es necesario
        max_length = 100000  # 100KB de texto
        if len(text) > max_length:
            text = text[:max_length] + "\n\n[CONTENIDO TRUNCADO]"
            logger.warning("Contenido truncado por exceder tamaño máximo")
        
        return text.strip()